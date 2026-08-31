from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(14)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return p

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Compass: DELSU Result Advisory System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Comprehensive Project Documentation\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    add_heading(doc, '1. Project Overview', 1)
    add_paragraph(doc, 'Compass is a comprehensive academic result tracking and AI advisory system built specifically for Delta State University (DELSU) Computer Science students. It bridges the gap caused by the university\'s often unavailable result portal by offering a robust dashboard where students can view their results, track their CGPA, identify carryovers, and interact with an intelligent AI Advisor.')
    
    add_heading(doc, '2. Core Features & Functionality', 1)
    
    add_heading(doc, '2.1. Student Dashboard & Features', 2)
    add_paragraph(doc, '• Instant Result Lookup: Students can log in and view their semester-by-semester academic performance without manual calculation.')
    add_paragraph(doc, '• Automated GPA & CGPA Tracking: The system automatically computes semester GPA and cumulative GPA based on uploaded broadsheet data.')
    add_paragraph(doc, '• Carryover Tracking: It identifies courses with failed grades (F) or missing prerequisites, presenting them clearly to the student.')
    add_paragraph(doc, '• At-Risk Flagging: Courses with "D" grades (scores between 45-49) are flagged as "At-Risk," providing early warnings for students who are barely passing.')
    
    add_heading(doc, '2.2. Adviser & Admin Flow', 2)
    add_paragraph(doc, '• Verified Access: Advisers must register and be explicitly approved by an Admin before they can access the upload system.')
    add_paragraph(doc, '• Automated Broadsheet Parsing: Advisers upload Excel (.xlsx) broadsheets. The backend (FastAPI + Pandas) parses the rows, extracting matriculation numbers, course codes, scores, and grades.')
    add_paragraph(doc, '• Secure Data Storage: Results are securely stored in a Supabase PostgreSQL database.')
    add_paragraph(doc, '• Analytics: Advisers have access to class performance metrics, top students, and at-risk students.')
    
    add_heading(doc, '2.3. The AI Academic Advisor', 2)
    add_paragraph(doc, '• Personalized Context: The AI Agent (powered by Groq/LLaMA) is injected with the specific student\'s entire academic history, current CGPA, total credits, and outstanding carryovers.')
    add_paragraph(doc, '• Predictive Reasoning: The AI is instructed to reason mathematically about the student\'s trajectory. If a student asks "What happens if I get an A in MTH 213?", the AI calculates the hypothetical impact on their CGPA.')
    add_paragraph(doc, '• Structured Output: Responses are strictly formatted in Markdown, avoiding generic greetings and prioritizing actionable academic advice.')
    
    add_heading(doc, '3. Technical Architecture', 1)
    
    add_heading(doc, '3.1. Frontend (React + Vite)', 2)
    add_paragraph(doc, '• Design System: A modern, mobile-first UI using Tailwind CSS. Core colors include midnight-ink, pure-canvas, mist, fog, and graphite.')
    add_paragraph(doc, '• State Management & Routing: Handled via React Context (useAuth) and React Router.')
    add_paragraph(doc, '• UI Components: Replaced native HTML tables with responsive, stacked cards for mobile devices. Interactive elements like the "ThinkingOrb" provide visual feedback during AI queries.')
    
    add_heading(doc, '3.2. Backend (FastAPI + Python)', 2)
    add_paragraph(doc, '• API Routes: Modularized routing (auth, students, upload, agent, analytics, results).')
    add_paragraph(doc, '• Database: Supabase (PostgreSQL) using the supabase-py client.')
    add_paragraph(doc, '• Data Processing: Pandas is used for efficient, robust Excel file parsing and data normalization.')
    add_paragraph(doc, '• LLM Integration: The Groq API handles fast, high-quality responses for the AI Advisor.')
    
    add_heading(doc, '4. Recent System Updates & Enhancements', 1)
    add_paragraph(doc, '• Mobile-First UI Overhaul: All complex table layouts (e.g., AdviserHistory) were refactored into stacked card layouts to eliminate horizontal scrolling on mobile.')
    add_paragraph(doc, '• Home Page Redesign: Restructured to tell a better story—highlighting the portal downtime problem, showcasing a 3-step "How It Works" flow, and featuring a dedicated mock-chat interface for the AI Advisor.')
    add_paragraph(doc, '• UI Standardization: Replaced all inconsistent padding and border-radii with the app\'s design tokens (rounded-[12px], text-step-sm, etc.).')
    add_paragraph(doc, '• Bug Fixes: Resolved an iOS Safari bug where the virtual keyboard caused the AI chat input box to float incorrectly in the middle of the screen by relying on native CSS `bottom-safe` positioning.')
    
    add_heading(doc, '5. Deployment', 1)
    add_paragraph(doc, '• Frontend: Deployed and automatically built via Vercel.')
    add_paragraph(doc, '• Backend: Designed for scalable deployment (e.g., Render, Railway, or Heroku) using Uvicorn.')
    
    # Save the document
    output_path = r'C:\Users\HP\Desktop\Compass_Project_Documentation.docx'
    doc.save(output_path)
    print(f"Documentation generated successfully at: {output_path}")

if __name__ == '__main__':
    main()
